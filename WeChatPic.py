import requests
import json
import time
from docx import Document
from docx.oxml.ns import qn
import os
from copy import copy
from urllib3 import encode_multipart_formdata

def clean_data():
    # 删除当前目录下的news.docx文件
    if os.path.exists("./news.docx"):
        os.remove("./news.docx")
        print("清除缓存中...")
    else:
        print("暂无缓存清除...")


def get_news():
    event_rate = []  # 事件评分
    event_name = []  # 事件名称
    event_desc = []  # 事件描述
    event_time = []  # 事件时间
    event_type = []  # 事件类型

    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    for i in range(1, 3): # 获取前两页数据
        print("正在获取第" + str(i) + "页数据")
        url = f"https://ef.zhiweidata.com/filterNew.do?firstType=全部&page={i}"
        headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.72 Safari/537.36"
        }
        
        response = requests.get(url=url, headers=headers)
        news = json.loads(response.text)
        for item in news["data"]["events"]:
            event_rate.append(item["index"])
            event_name.append(item["eventname"])
            event_desc.append(item["desc"])
            event_time.append(item["startTime"]/1000)
            event_type.append(item["type"])
    
    print("--------------------- 今日热点 ---------------------")
    for i in range(len(news["data"]["events"])):
        day_time = time.strftime('%Y-%m-%d', time.localtime(event_time[i]))
        # 判断日期是否是今天、昨天、前天
        if (day_time == time.strftime('%Y-%m-%d', time.localtime(time.time() - 86400*2)) or \
            day_time == time.strftime('%Y-%m-%d', time.localtime(time.time() - 86400)) or \
            day_time == time.strftime('%Y-%m-%d', time.localtime(time.time()))) and \
            event_type[i] != "娱乐":

            document.add_paragraph("影响力指数: " + str(event_rate[i]))
            document.add_paragraph(event_name[i])
            document.add_paragraph(event_desc[i])
            document.add_paragraph(day_time)
            document.add_paragraph("事件类型:" + event_type[i])
            document.add_paragraph("\n")
            print("影响力指数: ", event_rate[i])
            print(event_name[i])
            print(event_desc[i])
            print(time.strftime('%Y-%m-%d', time.localtime(event_time[i])))
            print("事件类型:", event_type[i])
            print("\n")
    document.save('news.docx')
    print("---------------------------------------------------")
    
def upload_file(file_path, wx_upload_url):
    file_name = file_path.split("/")[-1]
    with open(file_path, 'rb') as f:
        length = os.path.getsize(file_path)
        data = f.read()
    headers = {"Content-Type": "application/octet-stream"}
    params = {
        "filename": file_name,
        "filelength": length,
    }
    file_data = copy(params)
    file_data['file'] = (file_path.split('/')[-1:][0], data)
    encode_data = encode_multipart_formdata(file_data)
    file_data = encode_data[0]
    headers['Content-Type'] = encode_data[1]
    r = requests.post(wx_upload_url, data=file_data, headers=headers)
    print(r.text)
    media_id = r.json()['media_id']
    return media_id

# media_id 通过上一步上传的方法获得
def qi_ye_wei_xin_file(wx_url, media_id):
    headers = {"Content-Type": "text/plain"}
    data = {
        "msgtype": "file",
        "file": {
            "media_id": media_id
        }
    }
    r = requests.post(
        url=wx_url,
        headers=headers, json=data)
    print(r.text)


def push_report():
    test_report = './news.docx'
    wx_api_key = "xxxx" # Webkey
    wx_upload_url = "https://qyapi.weixin.qq.com/cgi-bin/webhook/upload_media?key={}&type=file".format(wx_api_key)
    wx_url = 'https://qyapi.weixin.qq.com/cgi-bin/webhook/send?key={}'.format(wx_api_key)
    media_id = upload_file(test_report, wx_upload_url)
    qi_ye_wei_xin_file(wx_url, media_id)
    print("------- 消息已经推送 -------")

if __name__ == '__main__':
    clean_data()
    time.sleep(1) # 防止文件未删除
    get_news()
    time.sleep(1) # 防止文件未生成
    push_report()
    
